from pathlib import Path
from pypdf import PdfReader
from docx import Document


def _extract_pdf(file_path: Path) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(file_path)
    text = ""

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text


def _extract_docx(file_path: Path) -> str:
    """Extract text from a DOCX file."""
    document = Document(file_path)
    text = ""

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text


def extract_text(file_path: Path) -> str:
    """Extract text from a supported resume file."""

    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return _extract_pdf(file_path)

    if extension == ".docx":
        return _extract_docx(file_path)

    raise ValueError(f"Unsupported file type: {extension}")